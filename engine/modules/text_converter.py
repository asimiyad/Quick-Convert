import os
import subprocess
import shutil
import codecs
from engine.core.base_engine import BaseConverter, validate_file, check_environment_dependencies
from engine.config import SUPPORTED_CONVERSIONS

try:
    import docx
except ImportError:
    docx = None

class TxtConverter(BaseConverter):
    """
    Handles converting Plain Text files (.txt) into other formats (e.g., DOCX, PDF).
    """

    def validate_requirements(self) -> bool:
        """
        Validates the text file size, configuration limits, and dependencies.
        """
        # 1. Base validation
        if not validate_file(self.input_file):
            return False

        # 2. Config validation
        _, input_ext = os.path.splitext(self.input_file)
        _, output_ext = os.path.splitext(self.output_file)
        input_ext, output_ext = input_ext.lower(), output_ext.lower()
        
        supported_outputs = SUPPORTED_CONVERSIONS.get(input_ext, [])
        if output_ext not in supported_outputs:
            self.logger.error(f"Conversion from {input_ext} to {output_ext} is disabled/unsupported by config.")
            return False

        # 3. Dynamic Dependencies
        if output_ext == ".docx" and docx is None:
            self.logger.error("Missing Python module 'python-docx'. Aborting. Install via: pip install python-docx")
            return False

        if output_ext == ".pdf":
            if not check_environment_dependencies(["soffice", "soffice.exe"]):
                 self.logger.warning("LibreOffice 'soffice' executable not cleanly mapped in system path.")

        self.logger.info("TXT format validation and dependency mapping physically successful.")
        return True

    def convert(self) -> bool:
        """
        Route to internal specific sub-workers depending on extension payload.
        """
        _, output_ext = os.path.splitext(self.output_file)
        output_ext = output_ext.lower()
        
        try:
            if output_ext == ".docx":
                return self._convert_to_docx()
            elif output_ext == ".pdf":
                 return self._convert_via_libreoffice("pdf")
            else:
                self.logger.error(f"Worker logic for {output_ext} is not explicitly implemented in TxtConverter.")
                return False
        except Exception as e:
            self.logger.exception(f"Critical error during physical TXT conversion loop: {e}")
            return False

    def _convert_to_docx(self) -> bool:
        """
        Dynamically reads lines of unstructured text and compiles them properly into a Microsoft Word object natively.
        """
        self.logger.debug(f"Constructing DOCX payload natively from raw text file: {self.input_file}")
        try:
            document = docx.Document()
            
            # Utilize fallback encoding to ignore harsh symbols that crash parsing visually
            with codecs.open(self.input_file, 'r', encoding='utf-8', errors='replace') as text_file:
                for line in text_file:
                    if line.strip():
                        document.add_paragraph(line.strip())
                    else:
                        # Retain spacing structure for empty lines natively
                        document.add_paragraph() 
                        
            document.save(self.output_file)
            self.logger.info(f"Successfully compiled clean structural DOCX payload natively: {self.output_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"DOCX compilation heavily failed under internal logic: {e}")
            return False

    def _convert_via_libreoffice(self, target_format: str) -> bool:
        """
        Constructs the output cleanly by bridging the engine to LibreOffice natively.
        """
        self.logger.info(f"Invoking headless LibreOffice to construct raw {target_format.upper()} payload.")
        outdir = os.path.dirname(self.output_file)
        
        command = [
            "soffice",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to", target_format,
            self.input_file,
            "--outdir", outdir
        ]
        
        try:
            result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            if result.returncode != 0:
                self.logger.error(f"LibreOffice backend violently failed:\n{result.stderr}")
                return False
                
            base_filename = os.path.basename(self.input_file)
            name_no_ext, _ = os.path.splitext(base_filename)
            lo_output_file = os.path.join(outdir, f"{name_no_ext}.{target_format}")
            
            if lo_output_file != self.output_file and os.path.exists(lo_output_file):
                shutil.move(lo_output_file, self.output_file)
                
            self.logger.info(f"Successfully constructed {target_format.upper()} document safely.")
            return True
            
        except FileNotFoundError:
             self.logger.error("CRITICAL: 'soffice' command utterly unresolvable logically via basic PATH mapping.")
             return False
        except Exception as e:
            self.logger.error(f"LibreOffice backend severely halted: {e}")
            return False
